#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模板解析器 — 解析任意 .docx 实验报告模板，自动识别结构和占位符
"""
import zipfile, json
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
from docx import Document


@dataclass
class CellInfo:
    row: int; col: int
    text: str
    is_merged: bool = False
    has_image: bool = False


@dataclass
class TableInfo:
    index: int
    rows: int; cols: int
    cells: List[CellInfo] = field(default_factory=list)
    merged_regions: List[Tuple] = field(default_factory=list)
    content_cells: List[CellInfo] = field(default_factory=list)   # 需要填充的单元格


@dataclass
class TemplateStructure:
    """解析后的模板结构"""
    filepath: str
    # 封页
    cover_paragraphs: List[dict] = field(default_factory=list)
    # 表格
    tables: List[TableInfo] = field(default_factory=list)
    # 占位符
    placeholders: Dict[str, str] = field(default_factory=dict)
    # 需要填充内容的关键单元格
    fillable_cells: List[dict] = field(default_factory=list)
    # 图片
    embedded_images: List[str] = field(default_factory=list)
    # 原始模板（用于复制）
    raw_bytes: bytes = b''


def parse_template(filepath: str) -> TemplateStructure:
    """解析模板的完整结构"""
    fp = Path(filepath)
    if not fp.exists():
        raise FileNotFoundError(f'模板文件不存在: {filepath}')

    ts = TemplateStructure(filepath=str(fp.absolute()))

    # 读取原始字节（用于复制）
    with open(filepath, 'rb') as f:
        ts.raw_bytes = f.read()

    # 提取嵌入图片
    with zipfile.ZipFile(filepath) as z:
        ts.embedded_images = [f for f in z.namelist() if 'media' in f.lower()]

    # 用 python-docx 解析
    doc = Document(filepath)

    # ── 解析封页段落 ──
    for p in doc.paragraphs:
        if p.text.strip():
            ts.cover_paragraphs.append({
                'text': p.text,
                'style': p.style.name,
                'alignment': str(p.alignment),
                'runs': [{'text': r.text, 'bold': r.bold, 'size': str(r.font.size)}
                         for r in p.runs],
            })

    # ── 解析表格结构 ──
    for ti, t in enumerate(doc.tables):
        tinfo = TableInfo(index=ti, rows=len(t.rows), cols=len(t.columns))

        # 检测合并单元格（内容相同的相邻单元格视为合并）
        cell_texts = {}
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                ctext = cell.text.strip()
                key = (ri, ci)

                # 检测图片
                has_img = False
                for p in cell.paragraphs:
                    for r in p.runs:
                        for child in r._element:
                            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                            if tag in ('drawing', 'pict'):
                                has_img = True

                info = CellInfo(ri, ci, ctext, has_image=has_img)
                tinfo.cells.append(info)
                cell_texts[key] = ctext

                # 判断是否可填充（内容非空且包含标题模式）
                if ctext and ('一、' in ctext or '二、' in ctext or '三、' in ctext
                              or '四、' in ctext or '五、' in ctext or '六、' in ctext):
                    tinfo.content_cells.append(info)

        # 检测合并区域
        for (r1, c1), t1 in cell_texts.items():
            if not t1: continue
            # 向右检查
            for c2 in range(c1 + 1, tinfo.cols):
                t2 = cell_texts.get((r1, c2), '')
                if t2 == t1:
                    tinfo.merged_regions.append((r1, c1, r1, c2))

        ts.tables.append(tinfo)

    # ── 识别占位符 ──
    # 从封页段落和表格标题行识别
    placeholder_patterns = {
        '学    院': 'college',
        '专    业': 'major',
        '班    级': 'class_name',
        '学生姓名': 'student_name',
        '学    号': 'student_id',
        '课程名称': 'course',
        '开课学期': 'semester',
        '实验日期': 'date',
        '实验项目名称': 'exp_title',
        '实验地点': 'location',
        '实验合作者姓名': 'partner',
    }

    all_text = ' '.join(p['text'] for p in ts.cover_paragraphs)
    for ti, tinfo in enumerate(ts.tables):
        for cell in tinfo.cells:
            all_text += ' ' + cell.text

    for pattern, key in placeholder_patterns.items():
        if pattern in all_text:
            ts.placeholders[key] = pattern

    # ── 识别需要填充内容的单元格 ──
    for ti, tinfo in enumerate(ts.tables):
        for ri in range(tinfo.rows):
            for ci in range(tinfo.cols):
                for cell in tinfo.cells:
                    if cell.row == ri and cell.col == ci:
                        # 检查是否是可填充的内容区
                        if any(marker in cell.text for marker in
                               ['一、实验目的', '二、实验内容', '三、实验原理',
                                '四、主要实验设备', '五、实验步骤', '六、实验结果',
                                '指导教师批阅', '总分']):
                            ts.fillable_cells.append({
                                'table_index': ti,
                                'row': ri,
                                'col': ci,
                                'hint': cell.text[:100],
                                'is_content': '实验目的' in cell.text,
                                'is_grade': '总分' in cell.text or '指导教师' in cell.text,
                            })
                            break

    return ts


def describe_for_llm(ts: TemplateStructure) -> str:
    """生成模板结构的 LLM 可读描述"""
    lines = []
    lines.append(f'## 实验报告模板结构分析')
    lines.append(f'模板文件: {Path(ts.filepath).name}')
    lines.append(f'嵌入图片: {len(ts.embedded_images)} 张')

    # 封页
    lines.append(f'\n### 封页信息')
    for p in ts.cover_paragraphs[:20]:
        lines.append(f'  {p["text"][:80]}')
    if len(ts.cover_paragraphs) > 20:
        lines.append(f'  ... (共 {len(ts.cover_paragraphs)} 段)')

    # 表格结构
    for t in ts.tables:
        lines.append(f'\n### 表格 {t.index+1} ({t.rows}行×{t.cols}列)')
        for cell in t.cells[:30]:
            img_tag = ' [含图片]' if cell.has_image else ''
            if cell.text.strip():
                lines.append(f'  [{cell.row},{cell.col}] {cell.text[:100]}{img_tag}')
        if len(t.cells) > 30:
            lines.append(f'  ... (共 {len(t.cells)} 个单元格)')

    # 占位符
    lines.append(f'\n### 识别到的占位符')
    for key, pattern in ts.placeholders.items():
        lines.append(f'  {key}: "{pattern}"')

    # 需要填充的单元格
    lines.append(f'\n### 需要生成内容的区域')
    for fc in ts.fillable_cells:
        lines.append(f'  表格{fc["table_index"]+1}[{fc["row"]},{fc["col"]}]: '
                     f'{fc["hint"][:80]}')

    return '\n'.join(lines)
