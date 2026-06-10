#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
工具链 — Word 文档组装、图片嵌入
"""

import os, shutil
from pathlib import Path
from typing import Dict, Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from template_parser import TemplateStructure


def set_cell_text(cell, text: str, font_name: str = '宋体',
                  size=Pt(10.5), color=None):
    """填充单元格文本内容"""
    # 清除原有段落内容
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
        p.clear()

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)

    lines = text.split('\n')
    for i, line in enumerate(lines):
        prefix = '\n' if i > 0 else ''
        run = p.add_run(prefix + line)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = size
        if color:
            run.font.color.rgb = color

    # 统一字体
    for run in p.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = size


def add_image_paragraph(cell, image_path: str, width_inches: float = 5.0,
                        caption: str = ''):
    """在单元格末尾添加图片段落"""
    if caption:
        cap_p = cell.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(caption)
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))

    return p


def add_mermaid_fallback(cell, mermaid_code: str):
    """当流程图 PNG 渲染失败时，嵌入 Mermaid 代码文本作为备用"""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    note = p.add_run('[软件流程图 — Mermaid 代码，可复制到 https://mermaid.live 查看]\n')
    note.font.size = Pt(8)
    note.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    code_p = cell.add_paragraph()
    code_run = code_p.add_run(mermaid_code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(7)
    code_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def assemble_report(
    template: TemplateStructure,
    content: Dict[str, str],
    flowchart_png_path: Optional[str],
    student_info: Dict[str, str],
    output_path: str,
    merge_placeholder: bool = True,
) -> str:
    """
    组装最终报告

    参数:
        template: 解析后的模板结构
        content: LLM 生成的内容 {title, purpose, content, principle, equipment, steps, results, mermaid}
        flowchart_png_path: 流程图 PNG 路径（可选）
        student_info: 学生信息
        output_path: 输出路径
        merge_placeholder: 是否替换模板中的占位符文本
    返回: 输出文件路径
    """
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    # 从原始模板字节创建副本
    with open(out, 'wb') as f:
        f.write(template.raw_bytes)

    doc = Document(str(out))

    # ── 填充表格内容 ──
    for ti, tinfo in enumerate(template.tables):
        if ti >= len(doc.tables):
            break
        t = doc.tables[ti]

        # --- 先填充标题行（Table 0, Row 0） ---
        if ti == 0:
            for ci in range(1, tinfo.cols):
                try:
                    tc = t.cell(0, ci)
                    if '实验项目名称' in t.cell(0, 0).text:
                        set_cell_text(tc, content.get('title', ''), size=Pt(10.5))
                except Exception:
                    pass

        # 查找并填充内容单元格
        for fc in template.fillable_cells:
            if fc['table_index'] != ti:
                continue
            try:
                cell = t.cell(fc['row'], fc['col'])

                # 根据 hint 判断填充什么内容
                hint = fc['hint']

                # 判断这是内容区（含一、二、三、四）还是步骤区（含五、六）
                if '一、' in hint and '二、' in hint:
                    # 四栏内容区
                    fill_text = (
                        f"一、实验目的\n{content.get('purpose', '')}\n\n"
                        f"二、实验内容\n{content.get('content', '')}\n\n"
                        f"三、实验原理与方法\n{content.get('principle', '')}"
                    )
                    set_cell_text(cell, fill_text)

                    # 嵌入流程图
                    if flowchart_png_path and os.path.exists(flowchart_png_path):
                        add_image_paragraph(cell, flowchart_png_path, 5.0,
                                           '图1 软件流程图')
                    elif content.get('mermaid', '').strip():
                        add_mermaid_fallback(cell, content['mermaid'])

                    # 设备器材
                    eq_text = f"\n四、主要实验设备及器材\n{content.get('equipment', '')}"
                    eq_p = cell.add_paragraph()
                    eq_run = eq_p.add_run(eq_text)
                    eq_run.font.name = '宋体'
                    eq_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    eq_run.font.size = Pt(10.5)

                elif '五、' in hint and '六、' in hint:
                    # 步骤+结果区
                    fill_text = (
                        f"五、实验步骤\n{content.get('steps', '')}\n\n"
                        f"六、实验结果与思考\n{content.get('results', '')}"
                    )
                    set_cell_text(cell, fill_text)

                elif '实验项目名称' in hint:
                    # 标题行
                    for ci in range(1, tinfo.cols):
                        try:
                            tc = t.cell(fc['row'], ci)
                            set_cell_text(tc, content.get('title', ''), size=Pt(10.5))
                        except Exception:
                            pass

            except Exception as e:
                print(f'    [警告] 填充 [{fc["row"]},{fc["col"]}] 失败: {e}')

    # ── 替换封页占位符（如果模板段落中有学生信息行）──
    if merge_placeholder:
        placeholder_map = {
            '学生姓名': student_info.get('name', ''),
            '学    号': student_info.get('id', ''),
            '班    级': student_info.get('class_name', ''),
            '专    业': student_info.get('major', ''),
            '学    院': student_info.get('college', ''),
            '课程名称': student_info.get('course', ''),
            '开课学期': student_info.get('semester', ''),
            '实验日期': student_info.get('date', ''),
            '实验地点': student_info.get('location', ''),
        }

        for p in doc.paragraphs:
            for key, value in placeholder_map.items():
                if key in p.text and value:
                    # 保留标签，追加内容
                    for run in p.runs:
                        if key in run.text or run.text.strip() in key:
                            # 在段落末尾追加值
                            pass
                    # 简化：直接在段落文本后添加
                    if not any(run.text.strip() == value for run in p.runs):
                        existing = p.text
                        if key in existing and value not in existing:
                            # 替换标签格式 "学生姓名：" → "学生姓名：张三"
                            new_text = existing.replace(key, f'{key}：{value}')
                            if new_text != existing:
                                p.clear()
                                run = p.add_run(new_text)
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.save(str(out))
    return str(out)
